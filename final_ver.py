import re
import pandas as pd
from pdfminer.high_level import extract_text
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx import Document  # from python-docx package

# === STEP 1: Extract text from PhonePe PDF ===
pdf_path = r"C:\VIGNESH REDDY\pp\PhonePe_Transaction_Statement_unlocked.pdf"
text = extract_text(pdf_path)

# Remove timestamps like "06:16 PM" (only if you don't need them)
# Comment this out if you want to use time for sorting later.
text = re.sub(r"\b\d{1,2}:\d{2}\s*(?:AM|PM)\b", "", text)

# === STEP 2: Regex pattern to extract transaction details ===
pattern = re.compile(
    r"([A-Za-z]{3} \d{2}, \d{4})[\s\S]*?(?:Paid to|Received from)\s+([A-Za-z0-9 ._@&\-*\n]+?)\s*(?=Transaction ID)"
    r"[\s\S]*?(Debit|Credit)\s*INR[\s\S]*?([\d,]+\.\d{2}|\d+)",
    re.DOTALL
)

transactions = []
for match in pattern.finditer(text):
    date = match.group(1).strip()
    name = re.sub(r"\s+", " ", match.group(2).strip())  # clean up names
    ttype = match.group(3).strip()
    amount = float(match.group(4).replace(",", ""))
    amount = -amount if ttype == "Debit" else amount
    transactions.append((date, name, amount, ttype))

if not transactions:
    print("⚠️ No transactions found. Check regex pattern.")
    exit()

# === STEP 3: Create DataFrame and sort ===
df = pd.DataFrame(transactions, columns=["Date", "Name", "Amount", "Type"])
df["Date"] = pd.to_datetime(df["Date"], format="%b %d, %Y")

# Preserve extraction order for intra-day sorting
df["Order"] = range(len(df))
df = df.sort_values(["Date", "Order"]).drop(columns=["Order"])

# === STEP 4: Build formatted text output ===
output_lines = []
for date, group in df.groupby("Date"):
    output_lines.append(f"{date.strftime('%b %d, %Y')}")
    for _, row in group.iterrows():
        sign = "+" if row["Amount"] > 0 else "-"
        output_lines.append(f"  {row['Name']:<45} {sign}{abs(row['Amount']):,.2f}")
    output_lines.append("")  # blank line between dates

formatted_text = "\n".join(output_lines)
print(formatted_text)

# === STEP 5: Export to Excel ===
df["Signed Amount"] = df["Amount"].apply(lambda x: f"{'+' if x > 0 else '-'}{abs(x):,.2f}")
df[["Date", "Name", "Type", "Signed Amount"]].to_excel("PhonePe_Transaction_Summary.xlsx", index=False)

# === STEP 6: Export to TXT (for debugging) ===
with open("PhonePe_Transaction_Summary.txt", "w", encoding="utf-8") as f:
    f.write(formatted_text)

# === STEP 7: Export to PDF ===
pdf_file = "PhonePe_Transaction_Summary.pdf"
doc = SimpleDocTemplate(pdf_file, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=50, bottomMargin=50)
styles = getSampleStyleSheet()
monospace = ParagraphStyle(
    'Monospace',
    parent=styles['Normal'],
    fontName='Courier',
    fontSize=10,
    leading=14
)
story = [Paragraph(line.replace(" ", "&nbsp;"), monospace) if line.strip() else Spacer(1, 0.1*inch)
         for line in output_lines]
doc.build(story)

# === STEP 8: Export to Word ===
word_file = "PhonePe_Transaction_Summary.docx"
docx = Document()
docx.add_heading("PhonePe Transaction Summary", level=1)
for line in output_lines:
    if line.strip():
        docx.add_paragraph(line)
    else:
        docx.add_paragraph("")  # blank line
docx.save(word_file)

# === STEP 9: Done ===
print("\n✅ Exported successfully:")
print("📘 Excel → PhonePe_Transaction_Summary.xlsx")
print("📄 PDF   → PhonePe_Transaction_Summary.pdf")
print("📝 Word  → PhonePe_Transaction_Summary.docx")
print("📜 Text  → PhonePe_Transaction_Summary.txt")
